import requests
from bs4 import BeautifulSoup
from docx import Document
import logging
import time

document = Document()
header = {'user-agent': 'curl/7.0'}
required_urls = []

def parse_page(url):
    review_html = requests.get(url, headers = header).text
    soup = BeautifulSoup(review_html, 'lxml')
    return soup

def get_urls():
    url = "https://paytm.com/"
    soup = parse_page(url)
    div_classes = ['_2kcr','_17WC']

    for div_class in div_classes:
        for data in soup.find_all('div', {'class': div_class}):
            for a in data.find_all('a'):
                required_urls.append(a.get('href'))

def get_data():
    div_sub_classes = ['nA7q', 'wk_w','_22U6']

    for url in required_urls:
        try:
            soup = parse_page(url)
            description_tag = soup.find(attrs={"name":"og:description"})
            description = description_tag["content"] if description_tag else "No Meta Description given"
            title_tag = soup.find("meta", property="og:title")            
            title = title_tag["content"] if title_tag else "No Meta Title given"
            for div_sub_class in div_sub_classes:
                for data in soup.find_all('div', {'class': div_sub_class}):
                    document.add_heading(url+ "\n")
                    document.add_paragraph("Title:\n "+ title)
                    document.add_paragraph("Description: \n "+ description)
                    document.add_paragraph("Data: \n "+ data.text)
        # Wait for 1 second
            time.sleep(1)
        except Exception as ex:            
            logging.exception("Exception Found in get_data method!")

def save_data():
    document.save('Paytm_Scraped.docx')

if __name__ == '__main__':
    get_urls()
    get_data()
    save_data()
